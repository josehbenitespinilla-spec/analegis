from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def limpiar(texto):
    if not texto:
        return ""
    return str(texto).strip()


def formato_ciudad(texto):
    if not texto:
        return ""
    return str(texto).strip().capitalize()


def configurar_documento(doc):

    section = doc.sections[0]

    section.page_width = Cm(21.59)
    section.page_height = Cm(33.02)

    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]

    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    pf = style.paragraph_format
    pf.line_spacing = 1
    pf.space_before = 0
    pf.space_after = 0


def agregar_parrafo(doc, texto=""):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def salto(doc, n=1):
    for _ in range(n):
        agregar_parrafo(doc, "")


def agregar_persona_parrafo(p, persona):

    tipo = persona.get("tipo", "natural")

    if tipo == "natural":

        ciudad = formato_ciudad(persona.get("ciudad"))

        r = p.add_run(limpiar(persona.get("nombre")).upper())
        r.bold = True

        p.add_run(", colombiano(a), mayor de edad, identificado(a) con cédula de ciudadanía No. ")
        p.add_run(limpiar(persona.get("identificacion")))
        p.add_run(" de ")
        p.add_run(ciudad)
        p.add_run(", domiciliado(a) en la ")
        p.add_run(limpiar(persona.get("direccion")))
        p.add_run(", correo electrónico ")
        p.add_run(limpiar(persona.get("correo")))
        p.add_run(", celular ")
        p.add_run(limpiar(persona.get("celular")))

    elif tipo == "juridica":

        r = p.add_run(limpiar(persona.get("razon_social")).upper())
        r.bold = True

        p.add_run(", identificada con NIT No. ")
        p.add_run(limpiar(persona.get("nit")))

        p.add_run(", representada legalmente por ")
        r = p.add_run(limpiar(persona.get("representante")).upper())
        r.bold = True

        p.add_run(", con domicilio en ")
        p.add_run(limpiar(persona.get("direccion")))

        p.add_run(", correo electrónico ")
        p.add_run(limpiar(persona.get("correo")))

        p.add_run(", celular ")
        p.add_run(limpiar(persona.get("celular")))

    elif tipo == "entidad":

        r = p.add_run(limpiar(persona.get("razon_social")).upper())
        r.bold = True

        p.add_run(", representada por ")
        r = p.add_run(limpiar(persona.get("representante")).upper())
        r.bold = True

        p.add_run(", con dirección ")
        p.add_run(limpiar(persona.get("direccion")))

        p.add_run(", correo electrónico ")
        p.add_run(limpiar(persona.get("correo")))

        p.add_run(", teléfono ")
        p.add_run(limpiar(persona.get("celular")))


def agregar_causantes(p, lista):

    lista = [c.strip().upper() for c in lista if c and c.strip() != ""]

    if not lista:
        return

    r = p.add_run(lista[0])
    r.bold = True

    for nombre in lista[1:]:
        p.add_run(" y ")
        r = p.add_run(nombre)
        r.bold = True


def generar(datos):

    doc = Document()
    configurar_documento(doc)

    ciudad = formato_ciudad(datos.get("ciudad"))
    fecha = limpiar(datos.get("fecha"))
    juzgado = limpiar(datos.get("juzgado")).upper()

    abogado_ciudad = formato_ciudad(datos.get("abogado_ciudad"))

    poderdantes = datos.get("poderdantes", [])
    demandados = datos.get("demandados", [])

    agregar_parrafo(doc, f"{ciudad}, {fecha}")

    salto(doc, 2)

    agregar_parrafo(doc, "Señores")

    p = agregar_parrafo(doc)
    p.add_run(juzgado).bold = True

    agregar_parrafo(doc, "E.     S.     D.")

    salto(doc, 2)

    p = agregar_parrafo(doc)
    p.add_run("Asunto: ").bold = True
    p.add_run("PODER ESPECIAL").bold = True

    salto(doc, 2)

    p = agregar_parrafo(doc)

    for i, persona in enumerate(poderdantes):

        if i > 0 and i == len(poderdantes) - 1:
            p.add_run(" y ")
        elif i > 0:
            p.add_run(", ")

        agregar_persona_parrafo(p, persona)

    p.add_run(", obrando en nombre propio otorgo ")
    p.add_run("PODER ESPECIAL AMPLIO Y SUFICIENTE").bold = True

    p.add_run(" al abogado(a) ")
    p.add_run(limpiar(datos.get("abogado_nombre")).upper()).bold = True

    p.add_run(", igualmente mayor y domiciliado(a) en la ")
    p.add_run(limpiar(datos.get("abogado_direccion")))

    p.add_run(", Abogado(a) Titulado(a) en ejercicio, identificado(a) con la Cédula de Ciudadanía No. ")
    p.add_run(limpiar(datos.get("abogado_identificacion")))

    p.add_run(", expedida en ")
    p.add_run(abogado_ciudad)

    p.add_run(", y portador(a) de la tarjeta profesional No. ")
    p.add_run(limpiar(datos.get("abogado_tarjeta")))

    p.add_run(" del C. S. de la Judicatura, correo electrónico ")
    p.add_run(limpiar(datos.get("abogado_correo")))

    p.add_run(", celular ")
    p.add_run(limpiar(datos.get("abogado_celular")))

    p.add_run(", para que en mi nombre solicite, tramite y lleve hasta su terminación proceso de ")
    p.add_run(limpiar(datos.get("tipo_proceso")).upper()).bold = True

    p.add_run(" en contra de ")

    for i, persona in enumerate(demandados):

        if i > 0 and i == len(demandados) - 1:
            p.add_run(" y ")
        elif i > 0:
            p.add_run(", ")

        agregar_persona_parrafo(p, persona)

    if datos.get("herederos_indeterminados"):

        lista = datos.get("causantes", [])

        if any(c and c.strip() for c in lista):
            p.add_run(" y contra los herederos indeterminados de ")
            agregar_causantes(p, lista)

    p.add_run(".")

    salto(doc, 1)

    agregar_parrafo(doc,
        "Mi apoderado cuenta con todas las facultades inherentes para el ejercicio "
        "del presente poder, en especial las de presentar demandas, reformarlas, "
        "desistir, sustituir, recibir, transigir, conciliar en diligencias "
        "prejudiciales y judiciales, recibir notificaciones, interponer recursos, "
        "solicitar pruebas, tacharlas, formular excepciones, reconvenir, llamar "
        "en garantía, denunciar el pleito, pactar cumplimiento, solicitar medidas "
        "cautelares y todas aquellas facultades consagradas en el artículo 77 del "
        "Código General del Proceso y demás normas concordantes, necesarias para "
        "la adecuada defensa de mis intereses."
    )

    salto(doc, 1)

    agregar_parrafo(doc,
        "Sírvase reconocerle personería en los términos y para los fines aquí señalados."
    )

    salto(doc, 1)

    agregar_parrafo(doc, "Atentamente,")

    salto(doc, 3)

    for persona in poderdantes:

        agregar_parrafo(doc, "____________________________")

        p = agregar_parrafo(doc)

        if persona.get("tipo") == "natural":
            p.add_run(limpiar(persona.get("nombre")).upper()).bold = True

            ciudad_p = formato_ciudad(persona.get("ciudad"))
            p = agregar_parrafo(doc)
            p.add_run(f"C.C. {limpiar(persona.get('identificacion'))} de {ciudad_p}").bold = True

        salto(doc, 2)

    p = agregar_parrafo(doc)
    p.add_run("Acepto").bold = True

    salto(doc, 3)

    agregar_parrafo(doc, "____________________________")

    p = agregar_parrafo(doc)
    p.add_run(limpiar(datos.get("abogado_nombre")).upper()).bold = True

    p = agregar_parrafo(doc)
    p.add_run(f"C.C. {limpiar(datos.get('abogado_identificacion'))} de {abogado_ciudad}").bold = True

    p = agregar_parrafo(doc)
    p.add_run(f"T.P. No. {limpiar(datos.get('abogado_tarjeta'))} del C. S. de la Judicatura.").bold = True

    return doc