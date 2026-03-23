from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def limpiar(t):
    return t.strip() if t else ""


def formato_ciudad(t):
    return t.strip().capitalize() if t else ""


def ciudad_departamento(ciudad, departamento):
    c = formato_ciudad(ciudad)
    d = limpiar(departamento)
    return f"{c}, {d}" if d else c


def configurar(doc):
    sec = doc.sections[0]

    sec.page_width = Cm(21.59)
    sec.page_height = Cm(33.02)

    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    pf = style.paragraph_format
    pf.line_spacing = 1
    pf.space_before = 0
    pf.space_after = 0


def p(doc, texto=""):
    par = doc.add_paragraph(texto)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


def salto(doc, n=1):
    for _ in range(n):
        p(doc, "")


def generar(datos):

    doc = Document()
    configurar(doc)

    ciudad = formato_ciudad(datos.get("ciudad"))
    fecha = limpiar(datos.get("fecha"))
    juzgado = limpiar(datos.get("juzgado")).upper()

    pod = datos["poderdante"]
    dem = datos["demandado"]

    # ENCABEZADO
    p(doc, f"{ciudad}, {fecha}")

    salto(doc, 2)

    p(doc, "Señores")

    pr = p(doc)
    pr.add_run(juzgado).bold = True

    p(doc, "E.     S.     D.")

    salto(doc, 2)

    pr = p(doc)
    pr.add_run("Asunto: ").bold = True
    pr.add_run("PODER ESPECIAL").bold = True

    salto(doc, 2)

    # CUERPO
    pr = p(doc)

    # PODERDANTE
    pr.add_run(limpiar(pod["nombre"]).upper()).bold = True
    pr.add_run(", colombiano(a), mayor de edad, identificado(a) con cédula de ciudadanía No. ")
    pr.add_run(limpiar(pod["identificacion"]))
    pr.add_run(", expedida en ")
    pr.add_run(ciudad_departamento(pod["ciudad"], pod.get("departamento")))
    pr.add_run(", domiciliado(a) en la ")
    pr.add_run(limpiar(pod["direccion"]))
    pr.add_run(", correo electrónico ")
    pr.add_run(limpiar(pod["correo"]))
    pr.add_run(", celular ")
    pr.add_run(limpiar(pod["celular"]))

    pr.add_run(", obrando en nombre propio otorgo ")
    pr.add_run("PODER ESPECIAL AMPLIO Y SUFICIENTE").bold = True

    # ABOGADO
    pr.add_run(" al abogado(a) ")
    pr.add_run(limpiar(datos["abogado_nombre"]).upper()).bold = True

    pr.add_run(", igualmente mayor, domiciliado(a) en la ")
    pr.add_run(limpiar(datos["abogado_direccion"]))

    pr.add_run(", Abogado(a) Titulado(a) en ejercicio, identificado(a) con la Cédula de Ciudadanía No. ")
    pr.add_run(limpiar(datos["abogado_identificacion"]))
    pr.add_run(", expedida en ")
    pr.add_run(ciudad_departamento(datos["abogado_ciudad"], datos.get("abogado_departamento")))

    pr.add_run(", y portador(a) de la tarjeta profesional No. ")
    pr.add_run(limpiar(datos["abogado_tarjeta"]))

    pr.add_run(" del C. S. de la Judicatura, correo electrónico ")
    pr.add_run(limpiar(datos["abogado_correo"]))

    pr.add_run(", celular ")
    pr.add_run(limpiar(datos["abogado_celular"]))

    pr.add_run(", para que en mi nombre solicite, tramite y lleve hasta su terminación proceso de ")
    pr.add_run(limpiar(datos["tipo_proceso"]).upper()).bold = True

    # DEMANDADO
    pr.add_run(" en contra de ")

    pr.add_run(limpiar(dem["nombre"]).upper()).bold = True
    pr.add_run(", colombiano(a), mayor de edad, identificado(a) con cédula de ciudadanía No. ")
    pr.add_run(limpiar(dem["identificacion"]))
    pr.add_run(", expedida en ")
    pr.add_run(ciudad_departamento(dem["ciudad"], dem.get("departamento")))
    pr.add_run(", domiciliado(a) en la ")
    pr.add_run(limpiar(dem["direccion"]))
    pr.add_run(", correo electrónico ")
    pr.add_run(limpiar(dem["correo"]))
    pr.add_run(", celular ")
    pr.add_run(limpiar(dem["celular"]))
    pr.add_run(".")

    salto(doc, 1)

    # PÁRRAFO LEGAL
    p(doc,
      "Mi apoderado cuenta con todas las facultades inherentes para el ejercicio del presente poder, "
      "en especial las de presentar demandas, reformarlas, desistir, sustituir, recibir, transigir, "
      "conciliar en diligencias prejudiciales y judiciales, recibir notificaciones, interponer recursos, "
      "solicitar pruebas, tacharlas, formular excepciones, reconvenir, llamar en garantía, denunciar el pleito, "
      "pactar cumplimiento, solicitar medidas cautelares y todas aquellas facultades consagradas en el artículo 77 "
      "del Código General del Proceso y demás normas concordantes, necesarias para la adecuada defensa de mis intereses."
    )

    salto(doc, 1)

    p(doc, "Sírvase reconocerle personería en los términos y para los fines aquí señalados.")

    salto(doc, 1)

    p(doc, "Atentamente,")

    salto(doc, 3)

    # FIRMA PODERDANTE
    p(doc, "____________________________")

    pr = p(doc)
    pr.add_run(limpiar(pod["nombre"]).upper()).bold = True

    pr = p(doc)
    pr.add_run(
        f"C.C. {limpiar(pod['identificacion'])} de {ciudad_departamento(pod['ciudad'], pod.get('departamento'))}"
    ).bold = True

    salto(doc, 2)

    # FIRMA ABOGADO
    pr = p(doc)
    pr.add_run("Acepto").bold = True

    salto(doc, 3)

    p(doc, "____________________________")

    pr = p(doc)
    pr.add_run(limpiar(datos["abogado_nombre"]).upper()).bold = True

    pr = p(doc)
    pr.add_run(
        f"C.C. {limpiar(datos['abogado_identificacion'])} de {ciudad_departamento(datos['abogado_ciudad'], datos.get('abogado_departamento'))}"
    ).bold = True

    pr = p(doc)
    pr.add_run(
        f"T.P. No. {limpiar(datos['abogado_tarjeta'])} del C. S. de la Judicatura."
    ).bold = True

    return doc