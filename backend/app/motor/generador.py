from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class GeneradorDocumento:

    def __init__(self):

        self.doc = Document()

        self.configurar_documento()

    # --------------------------------------------------
    # CONFIGURACION BASE DEL DOCUMENTO
    # --------------------------------------------------

    def configurar_documento(self):

        section = self.doc.sections[0]

        # TAMAÑO HOJA OFICIO
        section.page_height = Cm(33)
        section.page_width = Cm(21.6)

        # MARGENES 2.54 CM
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # ESTILO BASE
        style = self.doc.styles["Normal"]

        font = style.font
        font.name = "Arial"
        font.size = Pt(12)

        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    # --------------------------------------------------
    # NORMALIZAR CIUDAD (OPCIONAL)
    # --------------------------------------------------

    def capitalizar(self, texto):

        if not texto:
            return ""

        texto = texto.strip()

        return texto.capitalize()

    # --------------------------------------------------
    # NORMALIZAR NOMBRE PROPIO (OPCIONAL)
    # --------------------------------------------------

    def nombre_propio(self, texto):

        if not texto:
            return ""

        palabras = texto.split()

        palabras = [p.capitalize() for p in palabras]

        return " ".join(palabras)

    # --------------------------------------------------
    # CREAR PARRAFO JURIDICO
    # --------------------------------------------------

    def parrafo(self, texto, negrilla=False, cursiva=False):

        p = self.doc.add_paragraph()

        run = p.add_run(texto)

        run.bold = negrilla
        run.italic = cursiva

        run.font.name = "Arial"
        run.font.size = Pt(12)

        p_format = p.paragraph_format
        p_format.line_spacing = 1
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --------------------------------------------------
    # TITULO
    # --------------------------------------------------

    def titulo(self, texto):

        p = self.doc.add_paragraph()

        run = p.add_run(texto)

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_format = p.paragraph_format
        p_format.line_spacing = 1
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)

    # --------------------------------------------------
    # GUARDAR DOCUMENTO
    # --------------------------------------------------

    def guardar(self, ruta):

        self.doc.save(ruta)